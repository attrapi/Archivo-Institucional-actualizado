from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

GUINDA   = RGBColor(0x69, 0x1c, 0x32)
TERRACOTA= RGBColor(0xb8, 0x41, 0x0e)
GRIS     = RGBColor(0x4a, 0x40, 0x36)
BLANCO   = RGBColor(0xFF, 0xFF, 0xFF)

# ── Título principal ──
titulo = doc.add_heading('Dashboard · Archivo de Concentración ATTRAPI', level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.runs[0]
run.font.color.rgb = GUINDA
run.font.size = Pt(18)

# ── Subtítulo ──
sub = doc.add_paragraph('Descripción de pestañas del sistema de seguimiento documental')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = GRIS
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.italic = True

doc.add_paragraph()

# ── Introducción ──
intro = doc.add_paragraph(
    'El dashboard se alimenta en tiempo real desde un archivo de Google Sheets con 6 pestañas. '
    'Cada pestaña del sistema web corresponde directamente a una hoja del Excel. '
    'Cualquier cambio en el archivo se refleja automáticamente al recargar la página.'
)
intro.runs[0].font.size = Pt(11)
intro.runs[0].font.color.rgb = GRIS

doc.add_paragraph()

# ── Datos de cada pestaña ──
PESTANAS = [
    {
        'nombre': '1. Resumen',
        'hoja':   '— (combinación de todas las hojas)',
        'desc':   'Vista general del archivo. Muestra los indicadores clave (KPIs): total de cajas instaladas, legajos estimados, hojas estimadas, número de bodegas y unidades ATTRAPI. Incluye tres gráficas: distribución de cajas por bodega (dona), cajas por unidad administrativa (barras horizontales) y proyección acumulada 2025–2030 (línea).',
        'datos':  ['Total de cajas, legajos y hojas', 'Gráfica de distribución por bodega', 'Gráfica por unidad ATTRAPI', 'Línea de proyección sexenal'],
    },
    {
        'nombre': '2. Bodegas',
        'hoja':   '3_Estatus Archivo x Bodega',
        'desc':   'Inventario físico de los espacios de archivo. Muestra cada bodega (nombre, tipo de tenencia, superficie en m², cajas instaladas, legajos y hojas estimadas). Si se renombra una bodega en el Excel —por ejemplo de "Bodega J" a "Bodega Z"— el cambio aparece automáticamente en esta pestaña.',
        'datos':  ['Nombre y tipo de espacio (propio / prestado)', 'Superficie en m²', 'Cajas, legajos y hojas por bodega', '% del total de cajas'],
    },
    {
        'nombre': '3. Unidades',
        'hoja':   '1_Cajas insta of y ent',
        'desc':   'Concentración de cajas por unidad administrativa ATTRAPI. Muestra cuántas cajas llegaron por transferencia primaria y cuántas se heredaron por derogación de la DGDFM. Incluye gráfica de barras apiladas y tabla con buscador.',
        'datos':  ['Nombre de la unidad ATTRAPI', 'Cajas transferidas (transferencia primaria)', 'Cajas heredadas (derogación DGDFM)', 'Total por unidad'],
    },
    {
        'nombre': '4. Transferencias',
        'hoja':   '4_Transf Dir EXTINTAS',
        'desc':   'Registro detallado de cada movimiento de cajas desde las direcciones extintas. Permite filtrar por área generadora o bodega destino. Muestra dos gráficas: distribución por bodega de destino y evolución de transferencias por año.',
        'datos':  ['Área generadora', 'Número de cajas por movimiento', 'Bodega de destino (G, F, E, Z, etc.)', 'Fecha de la transferencia'],
    },
    {
        'nombre': '5. TP Mar 2026',
        'hoja':   '5._TP_ marzo_26',
        'desc':   'Transferencias primarias coordinadas por el Archivo de Concentración durante marzo 2026. Incluye unidades administrativas del ARTF y contratistas externos (Tren Interurbano, AIFA, etc.). Muestra el total acumulado de cajas y una tabla filtrable por unidad o descripción.',
        'datos':  ['Unidad administrativa o contratista', 'Número de transferencia (No. 01/2026, etc.)', 'Descripción del expediente', 'Número de cajas por transferencia'],
    },
    {
        'nombre': '6. Proyección',
        'hoja':   '2_Proyecc Sexenal',
        'desc':   'Estimación de crecimiento del acervo por obra ferroviaria de 2025 a 2030. Permite alternar entre vista acumulada y entradas anuales. Incluye gráfica de área total, gráfica de barras apiladas por obra y tabla con proyección por año.',
        'datos':  ['Obra ferroviaria (Tren Interurbano, Guadalajara, AIFA, etc.)', 'Existencia actual de cajas', 'Entradas anuales proyectadas por año', 'Acumulado proyectado hasta 2030'],
    },
]

for p in PESTANAS:
    # Encabezado de pestaña
    h = doc.add_heading(p['nombre'], level=1)
    h.runs[0].font.color.rgb = GUINDA
    h.runs[0].font.size = Pt(13)

    # Hoja de origen
    origen = doc.add_paragraph()
    r1 = origen.add_run('Hoja del Excel: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = GRIS
    r2 = origen.add_run(p['hoja'])
    r2.font.size = Pt(10)
    r2.font.color.rgb = TERRACOTA
    r2.italic = True

    # Descripción
    desc_p = doc.add_paragraph(p['desc'])
    desc_p.runs[0].font.size = Pt(10.5)
    desc_p.runs[0].font.color.rgb = GRIS

    # Datos que contiene
    etiq = doc.add_paragraph()
    r = etiq.add_run('Datos que contiene:')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GUINDA

    for dato in p['datos']:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.paragraph_format.left_indent = Cm(0.5)
        run = bullet.add_run(dato)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS

    doc.add_paragraph()

# ── Nota final ──
doc.add_paragraph()
linea = doc.add_paragraph('─' * 60)
linea.runs[0].font.color.rgb = RGBColor(0xd4, 0xca, 0xb3)

nota = doc.add_paragraph()
r = nota.add_run('Nota técnica: ')
r.bold = True
r.font.size = Pt(9.5)
r.font.color.rgb = GUINDA
r2 = nota.add_run(
    'El sistema no requiere actualizaciones manuales del código. '
    'Basta con editar el archivo en Google Sheets y recargar la página web para ver los cambios reflejados.'
)
r2.font.size = Pt(9.5)
r2.font.color.rgb = GRIS
r2.italic = True

# ── Pie ──
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp = pie.add_run('ATTRAPI · DGDFM · Archivo de Concentración · 2026')
rp.font.size = Pt(8.5)
rp.font.color.rgb = RGBColor(0x8a, 0x7f, 0x70)
rp.italic = True

doc.save('Descripcion_Dashboard_ATTRAPI.docx')
print("Archivo creado: Descripcion_Dashboard_ATTRAPI.docx")
