from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

GUINDA    = RGBColor(0x69, 0x1c, 0x32)
TERRACOTA = RGBColor(0xb8, 0x41, 0x0e)
AMBAR     = RGBColor(0xd7, 0xb2, 0x30)
VERDE     = RGBColor(0x6f, 0x9c, 0x58)
GRIS      = RGBColor(0x4a, 0x40, 0x36)
NEGRO     = RGBColor(0x1a, 0x16, 0x12)


def add_heading(text, size=14, color=GUINDA, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    h.runs[0].font.size = Pt(size)
    return h


def add_par(text, size=11, color=NEGRO, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_colored_bullet(color_label, color_rgb, description):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(color_label + ': ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = color_rgb
    r2 = p.add_run(description)
    r2.font.size = Pt(11)
    r2.font.color.rgb = NEGRO


# ── Título ──
titulo = doc.add_heading('Gráfica de ocupación y proyección por bodega', level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo.runs[0].font.color.rgb = GUINDA
titulo.runs[0].font.size = Pt(20)

sub = doc.add_paragraph('Pestaña "Área apilada" · Dashboard ATTRAPI')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = GRIS
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.italic = True

doc.add_paragraph()

# ── 1. Qué es ──
add_heading('1. ¿Qué tipo de gráfica es?')
add_par(
    'Es una gráfica de área apilada (en inglés, stacked area chart). Cada bodega '
    '(y el contenedor) aparece como un punto en el eje horizontal y sobre esa línea '
    'de puntos se trazan tres áreas de color que se apilan una sobre otra. La suma '
    'de las tres áreas en cada punto representa la superficie total relevante del '
    'espacio medida en metros cuadrados (m²). Este formato permite ver en un solo '
    'vistazo la composición interna de cada bodega y, al recorrer el gráfico de '
    'izquierda a derecha, comparar cómo se comporta el acervo entre espacios.'
)

# ── 2. Ejes ──
add_heading('2. ¿Qué representa cada eje?')
add_colored_bullet('Eje X (horizontal)', NEGRO,
    'nombre de cada bodega y del contenedor. El orden refleja cómo se cargaron los '
    'datos en la hoja de cálculo.')
add_colored_bullet('Eje Y (vertical)', NEGRO,
    'metros cuadrados (m²). La altura total de cada barra representa la superficie '
    'disponible del espacio más los m² que se proyecta ocupar hacia 2030.')

# ── 3. Colores ──
add_heading('3. ¿Qué significan los colores?')
add_colored_bullet('Rojo – Ocupado', TERRACOTA,
    'superficie que ya está siendo utilizada por las cajas existentes. Se calcula '
    'multiplicando el número de cajas de la bodega por la huella de una caja '
    '(0.35 m × 0.50 m = 0.175 m²).')
add_colored_bullet('Amarillo – Proyección', AMBAR,
    'metros cuadrados adicionales que se estima se ocuparán para el año 2030. Es la '
    'parte del crecimiento proyectado global que "le toca" a cada bodega según su '
    'participación actual en el total de cajas.')
add_colored_bullet('Verde – Disponible', VERDE,
    'superficie libre que aún queda en cada bodega al día de hoy (antes de aplicar '
    'la proyección). Cuando este segmento es pequeño o nulo, significa que la bodega '
    'está cerca o al tope de su capacidad.')

# ── 4. Cómo se calcula ──
add_heading('4. ¿Cómo se calculan los datos?')
add_par('Para cada bodega se realizan los siguientes cálculos:', bold=True)

p = doc.add_paragraph(style='List Number')
p.add_run('Ocupado (m²) = cajas actuales × 0.175').font.size = Pt(11)

p = doc.add_paragraph(style='List Number')
p.add_run('Disponible (m²) = máx(0, superficie del espacio − Ocupado)').font.size = Pt(11)

p = doc.add_paragraph(style='List Number')
run = p.add_run(
    'Proyección (m²) = (cajas proyectadas al 2030 − cajas existencia) × '
    '(cajas de esta bodega / total de cajas actuales) × 0.175'
)
run.font.size = Pt(11)

add_par(
    'La huella de 0.175 m² sale de las dimensiones físicas reales de una caja de '
    'archivo: 35 cm de ancho por 50 cm de largo. Los datos de existencia y de '
    'proyección 2030 provienen de la pestaña "Proyección" del mismo Google Sheet.',
    italic=True, color=GRIS, size=10
)

# ── 5. Lectura ──
add_heading('5. ¿Cómo se lee durante la exposición?')
add_par('Una barra alta indica un espacio con mucha superficie asignada en el '
        'acervo. Una barra en la que predomina el rojo es una bodega prácticamente '
        'saturada; cuando además tiene una franja amarilla grande, significa que '
        'recibirá todavía más carga hacia 2030 y representa un foco de riesgo.')
add_par('Si una barra tiene bastante verde y poca proyección amarilla, es una '
        'bodega con holgura y potencial de recibir nuevas transferencias. Si una '
        'barra queda casi toda amarilla por encima del rojo, la ocupación proyectada '
        'superaría la capacidad disponible y sería necesario redistribuir o habilitar '
        'un nuevo espacio.')

# ── 6. Valores en pantalla ──
add_heading('6. Valores que aparecen en pantalla')
add_par('Sobre el pico de cada bodega se muestra una etiqueta blanca con el total '
        'en m² (Ocupado + Proyección + Disponible). Debajo de la gráfica hay una '
        'franja de tarjetas con el desglose por bodega: el nombre del espacio, una '
        'mini-barra de proporciones con los tres colores y el valor numérico exacto '
        'de Ocupado, Proyección y Disponible. Si se pasa el cursor sobre la gráfica '
        'aparece un tooltip con el detalle, la superficie del espacio y el número '
        'de cajas actuales.')

# ── 7. Puntos clave ──
add_heading('7. Puntos clave para la presentación')
bullets = [
    'La gráfica responde a la pregunta: "¿cuánto espacio ocupa hoy cada bodega, '
    'cuánto queda libre y cuánto más ocupará en el próximo sexenio?"',
    'Los tres colores (rojo, amarillo, verde) funcionan como un semáforo visual: '
    'rojo = lo que ya está consumido; amarillo = lo que viene; verde = lo que aún '
    'hay disponible.',
    'La proyección no se reparte en partes iguales entre bodegas: se pondera según '
    'qué tanto representa cada bodega del total actual de cajas, asumiendo que los '
    'patrones de transferencia se mantienen.',
    'Si la suma del rojo y el amarillo excede el total histórico de la bodega, '
    'eso es una alerta operativa: la capacidad actual no alcanzará.',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(b)
    r.font.size = Pt(11)
    r.font.color.rgb = NEGRO

doc.add_paragraph()
cierre = doc.add_paragraph('Fuente: Dashboard Archivo de Concentración · ATTRAPI')
cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
cierre.runs[0].font.size = Pt(9)
cierre.runs[0].font.italic = True
cierre.runs[0].font.color.rgb = GRIS

doc.save('Grafica_Area_Apilada_ATTRAPI.docx')
print('Documento generado: Grafica_Area_Apilada_ATTRAPI.docx')
